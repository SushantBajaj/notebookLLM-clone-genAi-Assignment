from __future__ import annotations

import csv
import re
from pathlib import Path


class DocumentParseError(Exception):
    pass


def parse_document(file_path: str | Path, extension: str) -> str:
    path = Path(file_path)
    parser = {
        "csv": _parse_csv,
        "doc": _parse_doc,
        "docx": _parse_docx,
        "pdf": _parse_pdf,
    }.get(extension.lower())

    if parser is None:
        raise DocumentParseError(f"Unsupported file type: {extension}")

    text = parser(path).strip()
    if not text:
        raise DocumentParseError("No readable text found in document.")

    return text


def _parse_csv(path: Path) -> str:
    rows: list[str] = []

    with path.open("r", encoding="utf-8-sig", newline="") as csv_file:
        reader = csv.reader(csv_file)
        for row in reader:
            rows.append(" | ".join(cell.strip() for cell in row))

    return "\n".join(rows)


def _parse_doc(path: Path) -> str:
    raw = path.read_bytes()

    for encoding in ("utf-8", "latin-1"):
        try:
            text = raw.decode(encoding, errors="ignore")
            return _clean_text(text)
        except UnicodeDecodeError:
            continue

    raise DocumentParseError("Could not read legacy DOC file.")


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError("Install python-docx to parse DOCX files.") from exc

    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_rows = []

    for table in document.tables:
        for row in table.rows:
            table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))

    return "\n".join([*paragraphs, *table_rows])


def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError("Install pypdf to parse PDF files.") from exc

    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _clean_text(text: str) -> str:
    printable = re.sub(r"[^\x09\x0a\x0d\x20-\x7e]+", " ", text)
    return re.sub(r"[ \t]+", " ", printable)
